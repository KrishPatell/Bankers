"""Safely add one DOCX-authored post to the Bankers Vascular Blogs export.

This is a local publishing-preparation tool: it changes the CMS CSV and copies
the supplied image into ``src/images/blog/``.  It deliberately never builds,
deploys, commits, or overwrites an existing article.
"""

import argparse
import csv
import html
import os
import re
import shutil
import sys
import unicodedata
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
CMS = ROOT / "cms"
BLOGS = next(CMS.glob("*- Blogs - *.csv"))
AUTHORS = next(CMS.glob("*- Blog Authors - *.csv"))
DOCTORS = next(CMS.glob("*- Our Doctors - *.csv"))
IMAGE_DIR = ROOT / "src" / "images" / "blog"


def normalise(value):
    value = unicodedata.normalize("NFKD", value or "")
    value = "".join(c for c in value if not unicodedata.combining(c)).lower()
    return re.sub(r"[^a-z0-9]+", "", value.replace("doctor", "dr"))


def slugify(title):
    text = unicodedata.normalize("NFKD", title).encode("ascii", "ignore").decode()
    text = re.sub(r"[^a-z0-9]+", "-", text.lower()).strip("-")
    return text[:110].rstrip("-")


def read_csv(path):
    with open(path, encoding="utf-8-sig", newline="") as fh:
        reader = csv.DictReader(fh)
        return reader.fieldnames or [], list(reader)


def safe_published(rows):
    return [r for r in rows if (r.get("Archived") or "").lower() != "true"
            and (r.get("Draft") or "").lower() != "true"]


def select_author(name):
    """Prefer the existing Blog Author record; otherwise use a live doctor.

    A doctor-only match uses the existing /our-doctors profile as the author
    URL.  No author or doctor record is created by this workflow.
    """
    wanted = normalise(name)
    for path, label in ((AUTHORS, "blog-author"), (DOCTORS, "our-doctors")):
        _, rows = read_csv(path)
        matches = [r for r in safe_published(rows) if normalise(r.get("Name")) == wanted]
        if len(matches) == 1:
            return matches[0]["Slug"].strip(), matches[0]["Name"].strip(), label
        if len(matches) > 1:
            raise ValueError("more than one published %s profile matches %r" % (label, name))
    raise ValueError("no published Blog Author or doctor profile matches %r" % name)


def run_html(paragraph):
    """Use run formatting when Word exposes it; retain all text as a fallback."""
    text = paragraph.text
    if not text.strip():
        return ""
    if "".join(r.text for r in paragraph.runs) != text:
        return html.escape(text)
    out = []
    for run in paragraph.runs:
        value = html.escape(run.text).replace("\n", "<br>")
        if run.bold:
            value = "<strong>%s</strong>" % value
        if run.italic:
            value = "<em>%s</em>" % value
        if run.underline:
            value = "<u>%s</u>" % value
        out.append(value)
    return "".join(out)


def paragraph_image_rel_ids(paragraph):
    """Return inline-image relationship IDs in their authored paragraph order."""
    return re.findall(r'r:embed="(rId\d+)"', paragraph._p.xml)


def paragraph_kind(paragraph):
    style = (paragraph.style.name if paragraph.style else "").lower()
    if "heading 1" in style or style == "title":
        return "h1"
    if "heading 2" in style:
        return "h2"
    if "heading 3" in style:
        return "h3"
    if "list bullet" in style or "bullet" in style:
        return "ul"
    if "list number" in style or "number" in style:
        return "ol"
    return "p"


def inferred_list_items(items):
    """Recognise simple DOCX lists pasted as short normal paragraphs.

    Clinicians commonly paste a bullet list from another editor without its
    Word list style. We retain its text but render a real semantic list when a
    run of short, non-sentence paragraphs clearly behaves like one.
    """
    flags, index = set(), 0
    while index < len(items):
        item = items[index]
        if not isinstance(item, Paragraph) or paragraph_kind(item) != "p":
            index += 1
            continue
        run, cursor = [], index
        while cursor < len(items):
            candidate = items[cursor]
            if not isinstance(candidate, Paragraph) or paragraph_kind(candidate) != "p":
                break
            text = candidate.text.strip()
            if not text or len(text) > 115 or text.endswith((".", "?", "!", ":")):
                break
            run.append(cursor)
            cursor += 1
        # Two items after a colon, or three consecutive terse items, is a
        # sufficiently strong signal without converting ordinary prose.
        lead = items[index - 1] if index else None
        lead_text = lead.text.strip() if isinstance(lead, Paragraph) else ""
        if len(run) >= 3 or (len(run) >= 2 and lead_text.endswith(":")):
            flags.update(run)
        index = cursor if cursor > index else index + 1
    return flags


def docx_to_html(path):
    document = Document(path)
    # ``document.paragraphs`` and ``document.tables`` are separate lists. Walk
    # the OOXML body so tables remain in their authored position in the article.
    items = []
    for child in document.element.body.iterchildren():
        if child.tag.endswith("}p"):
            items.append(Paragraph(child, document))
        elif child.tag.endswith("}tbl"):
            items.append(Table(child, document))
    paras = [p for p in items if isinstance(p, Paragraph) and p.text.strip()]
    if not paras:
        raise ValueError("the DOCX contains no readable text")
    title = paras[0].text.strip()
    blocks, list_kind, list_items = [], None, []

    def flush_list():
        nonlocal list_kind, list_items
        if list_items:
            blocks.append("<%s>%s</%s>" % (list_kind, "".join(list_items), list_kind))
        list_kind, list_items = None, []

    title_paragraph = paras[0]
    plain_list_indexes = inferred_list_items(items)
    for item_index, item in enumerate(items):
        if isinstance(item, Table):
            flush_list()
            rows = []
            for r_index, row in enumerate(item.rows):
                cells = [html.escape(cell.text.strip()).replace("\n", "<br>") for cell in row.cells]
                tag = "th" if r_index == 0 else "td"
                rows.append("<tr>%s</tr>" % "".join("<%s>%s</%s>" % (tag, cell, tag)
                                                       for cell in cells))
            if rows:
                blocks.append("<table><tbody>%s</tbody></table>" % "".join(rows))
            continue
        paragraph = item
        image_ids = paragraph_image_rel_ids(paragraph)
        if not paragraph.text.strip() and not image_ids:
            continue
        # The CMS detail template renders the title as the page's one H1. Do
        # not repeat it in rich text; it remains the visible page title.
        if paragraph is title_paragraph:
            continue
        kind, body = paragraph_kind(paragraph), run_html(paragraph)
        # Word may place an image in a paragraph with no text. Keep it at the
        # exact authored point instead of dropping it during text conversion.
        if image_ids:
            flush_list()
            blocks.extend('<figure class="blog-inline-image" data-docx-image="%s"></figure>' % image_id
                          for image_id in image_ids)
        if item_index in plain_list_indexes:
            kind = "ul"
        if not body:
            continue
        if kind in ("ul", "ol"):
            if list_kind != kind:
                flush_list()
                list_kind = kind
            list_items.append("<li>%s</li>" % body)
            continue
        flush_list()
        # The template owns the only H1. Any later Word Heading 1 is retained
        # as a lower-level web heading, never discarded.
        if kind == "h1":
            kind = "h2"
        blocks.append("<%s>%s</%s>" % (kind, body, kind))

    flush_list()
    return title, "\n".join(blocks)


def extract_docx_images(path, slug):
    """Copy original inline DOCX images and return their local public URLs.

    The blog workflow retains clinician-supplied inline visuals rather than
    silently publishing a text-only article. Names are deterministic and are
    never allowed to overwrite an existing blog asset.
    """
    document = Document(path)
    urls, written = {}, []
    try:
        for paragraph in document.paragraphs:
            for image_id in paragraph_image_rel_ids(paragraph):
                if image_id in urls:
                    continue
                part = document.part.related_parts.get(image_id)
                if not part or not getattr(part, "blob", None):
                    continue
                ext = {
                    "image/png": ".png", "image/jpeg": ".jpg", "image/gif": ".gif",
                    "image/webp": ".webp", "image/tiff": ".tiff",
                }.get(getattr(part, "content_type", ""), ".png")
                dest = IMAGE_DIR / ("%s-inline-%d%s" % (slug, len(urls) + 1, ext))
                if dest.exists():
                    if dest.read_bytes() != part.blob:
                        raise ValueError("refusing to overwrite a different inline image: %s" % dest)
                    urls[image_id] = "/images/blog/%s" % dest.name
                    continue
                dest.write_bytes(part.blob)
                written.append(dest)
                urls[image_id] = "/images/blog/%s" % dest.name
        return urls, written
    except Exception:
        for image in written:
            image.unlink(missing_ok=True)
        raise


TOPIC_TARGETS = (
    (r"\b(heel pain|plantar fasciitis|heel embolization)\b", "heel pain", "departments", "heel-pain"),
    (r"\b(knee|osteoarthritis|gae|genicular)\b", "knee pain", "treatment", "genicular"),
    (r"\b(varicose|vein|venous|sclerotherapy)\b", "varicose veins", "treatment", "scler"),
    (r"\b(prostate|pae)\b", "prostate", "treatment", "prostate"),
    (r"\b(piles|hemorrhoid|hae)\b", "piles", "treatment", "hemorrhoid"),
    (r"\b(hair|prp)\b", "hair", "departments", "hair"),
)

# Related posts are linked only when their topic phrase naturally occurs in
# the supplied article and the target slug is currently published.
RELATED_BLOG_TARGETS = (
    (r"\bheel pain embolization\b", "Heel Pain Embolization",
     "how-embolization-helps-in-chronic-heel-pain"),
)


def collection_rows(label):
    label = {"blog": "Blogs", "treatment": "Treatments", "departments": "Departments"}[label]
    path = next(CMS.glob("*- %s - *.csv" % label))
    _, rows = read_csv(path)
    return safe_published(rows)


def find_target(collection, token):
    for row in collection_rows(collection):
        if token in (row.get("Name") or "").lower() or token in (row.get("Slug") or "").lower():
            folder = {"treatment": "treatment", "departments": "departments"}[collection]
            return "/%s/%s" % (folder, row["Slug"].strip()), row["Name"].strip()
    return None


def find_published_blog(slug):
    for row in collection_rows("blog"):
        if (row.get("Slug") or "").strip() == slug:
            return "/blog/%s" % slug
    return None


def link_first_text(html_body, phrase, url):
    """Add one natural link, never inside headings, links, or table markup."""
    pattern = re.compile(r"(<(?:p|li)[^>]*>)(.*?)(</(?:p|li)>)", re.I | re.S)
    phrase_re = re.compile(r"(?<![\w-])(%s)(?![\w-])" % re.escape(phrase), re.I)

    def replace_block(match):
        nonlocal changed
        if changed or "<a " in match.group(2).lower():
            return match.group(0)
        value, count = phrase_re.subn(r'<a href="%s">\1</a>' % url, match.group(2), count=1)
        if count:
            changed = True
        return match.group(1) + value + match.group(3)

    changed = False
    return pattern.sub(replace_block, html_body)


def add_internal_links(body):
    """Link only phrases actually present in prose; no appended SEO filler."""
    plain = re.sub(r"<[^>]+>", " ", body).lower()
    linked = set()
    for topic_re, natural_phrase, collection, token in TOPIC_TARGETS:
        if not re.search(topic_re, plain):
            continue
        found = find_target(collection, token)
        if not found or found[0] in linked:
            continue
        url, name = found
        for phrase in sorted({natural_phrase, name, token, token.title()}, key=len, reverse=True):
            amended = link_first_text(body, phrase, url)
            if amended != body:
                body, linked = amended, linked | {url}
                break
    # Existing blog titles are linked only when that exact title is already a
    # natural part of the supplied prose.  This avoids bolting a "related"
    # sentence onto the clinician's content solely for SEO.
    for row in collection_rows("blog"):
        title = (row.get("Name") or "").strip()
        if len(title) < 12:
            continue
        url = "/blog/%s" % (row.get("Slug") or "").strip()
        if not url.endswith("/") and url not in linked:
            amended = link_first_text(body, title, url)
            if amended != body:
                body, linked = amended, linked | {url}
    for topic_re, phrase, slug in RELATED_BLOG_TARGETS:
        url = find_published_blog(slug)
        if url and re.search(topic_re, plain) and url not in linked:
            amended = link_first_text(body, phrase, url)
            if amended != body:
                body, linked = amended, linked | {url}
    return body


def short_description(body, title):
    # Prefer the first prose paragraph over an H1/H2 title or a Word section
    # label such as "Introduction". This leaves all article content untouched
    # while producing a useful search/social excerpt.
    first_paragraph = re.search(r"<p[^>]*>(.*?)</p>", body, re.I | re.S)
    text = first_paragraph.group(1) if first_paragraph else body
    text = re.sub(r"<[^>]+>", " ", text)
    text = re.sub(r"\s+", " ", html.unescape(text)).strip()
    if text.lower().startswith(title.lower()):
        text = text[len(title):].strip()
    return (text or title)[:155].rstrip(" ,;:-")


def write_row(fields, rows, row):
    for key in row:
        if key not in fields:
            fields.append(key)
    with open(BLOGS, "w", encoding="utf-8", newline="") as fh:
        writer = csv.DictWriter(fh, fieldnames=fields, extrasaction="ignore")
        writer.writeheader()
        # A retry after a transient build/publish interruption must update the
        # prepared row, not create a duplicate public URL.
        writer.writerows([existing for existing in rows
                          if (existing.get("Slug") or "").strip().lower() != row["Slug"].lower()] + [row])


def main():
    parser = argparse.ArgumentParser(description="Prepare one Bankers Vascular blog from DOCX + thumbnail.")
    parser.add_argument("docx", type=Path)
    parser.add_argument("thumbnail", type=Path)
    author_group = parser.add_mutually_exclusive_group(required=True)
    author_group.add_argument("--author", help="Existing doctor or Blog Author name")
    author_group.add_argument("--bankers-notes", action="store_true",
                              help="Publish as a Dr. Mohal Banker Bankers Note")
    parser.add_argument("--published-on", type=lambda value: datetime.strptime(value, "%Y-%m-%d"),
                        help="Scheduled publication date in YYYY-MM-DD (defaults to today)")
    parser.add_argument("--listing-priority", type=int, default=0,
                        help="Optional blog listing priority; higher values appear before date-sorted posts")
    parser.add_argument("--dry-run", action="store_true", help="Validate and preview without writing files")
    args = parser.parse_args()
    if not args.docx.is_file() or args.docx.suffix.lower() != ".docx":
        raise SystemExit("DOCX file not found or not a .docx: %s" % args.docx)
    if not args.thumbnail.is_file():
        raise SystemExit("thumbnail file not found: %s" % args.thumbnail)

    title, rich = docx_to_html(args.docx)
    fields, rows = read_csv(BLOGS)
    slug = slugify(title)
    if not slug:
        raise SystemExit("could not derive a URL slug from the DOCX title")
    existing_same_slug = None
    for existing in rows:
        if normalise(existing.get("Name")) == normalise(title):
            existing_same_slug = existing
        if (existing.get("Slug") or "").strip().lower() == slug.lower():
            existing_same_slug = existing
    if existing_same_slug and not existing_same_slug.get("Slug"):
        raise SystemExit("duplicate blog title: %s" % existing_same_slug.get("Name"))

    requested_author = "Dr. Mohal Banker" if args.bankers_notes else args.author
    author_slug, author_name, author_kind = select_author(requested_author)
    rich = add_internal_links(rich)
    extension = args.thumbnail.suffix.lower()
    if extension not in {".avif", ".gif", ".jpeg", ".jpg", ".png", ".webp"}:
        raise SystemExit("thumbnail must be AVIF, GIF, JPG, PNG, or WebP")
    image_rel = "/images/blog/%s%s" % (slug, extension)
    published_at = args.published_on.replace(tzinfo=timezone.utc) if args.published_on else datetime.now(timezone.utc)
    now = published_at.strftime("%a %b %d %Y 00:00:00 GMT+0000 (Coordinated Universal Time)")
    row = {field: "" for field in fields}
    row.update({
        "Name": title, "Slug": slug, "Archived": "false", "Draft": "false",
        "Created On": now, "Updated On": now, "Published On": now, "time": now,
        "Blog Thumbnail": image_rel, "Main Image": image_rel,
        "Short Details": short_description(rich, title), "Main Details": rich,
        # The build appends the established brand suffix exactly once.
        "Author": author_slug, "Meta Title": title,
        "Meta Description": short_description(rich, title),
        "Listing Priority": str(args.listing_priority),
    })
    print("Title: %s\nSlug: %s\nAuthor: %s (%s)\nCanonical: https://bankersvascular.com/blog/%s"
          % (title, slug, author_name, author_kind, slug))
    if args.dry_run:
        return
    IMAGE_DIR.mkdir(parents=True, exist_ok=True)
    image_dest = IMAGE_DIR / (slug + extension)
    if not image_dest.exists():
        shutil.copy2(args.thumbnail, image_dest)
    try:
        inline_urls, inline_images = extract_docx_images(args.docx, slug)
        for image_id, image_url in inline_urls.items():
            rich = rich.replace(
                '<figure class="blog-inline-image" data-docx-image="%s"></figure>' % image_id,
                '<figure class="blog-inline-image"><img src="%s" alt="%s"></figure>'
                % (image_url, html.escape(title)),
            )
        row["Main Details"] = rich
        write_row(fields, rows, row)
    except Exception:
        image_dest.unlink(missing_ok=True)
        for inline_image in locals().get("inline_images", []):
            inline_image.unlink(missing_ok=True)
        raise
    print("Prepared CMS row and copied thumbnail. Run the build and verification before publishing.")


if __name__ == "__main__":
    main()
